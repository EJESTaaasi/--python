import os
import logging
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium import webdriver
from docx import Document
from docx.shared import Inches
import pyautogui
from datetime import datetime

# 配置日志
logging.basicConfig(level=logging.INFO)

# 根据当前时间创建文件夹
def create_folder_for_current_time(base_dir, current_time):
    dir_path = os.path.join(base_dir, current_time)
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    return dir_path

# 获取当前时间
current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

# 设置截图和文档的存储目录
screenshot_dir = create_folder_for_current_time("screenshots", current_time)
document_dir = create_folder_for_current_time("documents", current_time)

# 初始化Chrome浏览器驱动
chrome_options = Options()
chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
chrome_options.add_experimental_option('useAutomationExtension', False)
chrome_service = Service(executable_path=r'C:\Program Files\Google\Chrome\Application\chromedriver.exe')  # 确保这是您ChromeDriver的实际路径
driver = webdriver.Chrome(service=chrome_service, options=chrome_options)

# 最大化浏览器窗口以捕捉更多内容
driver.maximize_window()

# 创建Word文档
doc = Document()

try:
    with open('urls.txt', 'r') as file:
        urls = [line.strip() for line in file]

    # 遍历URLs，打开并截图整个桌面
    for index, url in enumerate(urls):
        try:
            logging.info(f"Visiting {url}")
            driver.get(url)
            driver.implicitly_wait(10)  # 等待页面加载完成

            logging.info(f"Capturing the entire desktop, including the time on the taskbar for: {url}")

            # 截图整个桌面
            screenshot_path = os.path.join(screenshot_dir, f'screenshot_{index}.png')
            pyautogui.screenshot(screenshot_path)

            logging.info(f"Saved screenshot to {screenshot_path}")

            # 将截图插入Word文档
            doc.add_picture(screenshot_path, width=Inches(6))
            doc.add_paragraph(url)  # 在截图下方添加URL
            doc.add_page_break()  # 每个截图后添加分页

        except Exception as e:
            logging.error(f"Error when visiting {url}: {e}")

finally:
    # 保存Word文档
    word_filename = os.path.join(document_dir, f'Screenshots_Document_{current_time}.docx')
    doc.save(word_filename)
    logging.info(f"Word document saved to {word_filename}")
    
    # 关闭浏览器
    driver.quit()
